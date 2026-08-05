"""
Resume Generator - Create .docx output from tailored resume content

Converts tailored resume data into ATS-compliant .docx format.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import os

def generate_docx_resume(resume_data, output_dir=None):
    """
    Generate a .docx file from tailored resume data.
    
    Args:
        resume_data (dict): Tailored resume content with sections:
            - name (str): Candidate name
            - email (str): Contact email
            - phone (str): Contact phone
            - summary (str): Professional summary
            - skills (list): Skills list
            - experience (list): Work experience entries
            - education (list): Education entries
            - certifications (list): Optional certifications
        output_dir (str): Directory to save .docx file. Defaults to project root.
    
    Returns:
        dict: Status with file path and confirmation
    """
    try:
        # Create new Document
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Header: Name and Contact Info
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(resume_data.get('name', 'Your Name'))
        name_run.font.size = Pt(16)
        name_run.font.bold = True
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        contact_info = f"{resume_data.get('email', '')} | {resume_data.get('phone', '')}"
        contact_para = doc.add_paragraph(contact_info)
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.runs[0].font.size = Pt(10)
        
        doc.add_paragraph()  # Spacing
        
        # Professional Summary
        if resume_data.get('summary'):
            doc.add_heading('PROFESSIONAL SUMMARY', level=2)
            doc.add_paragraph(resume_data['summary'])
            doc.add_paragraph()
        
        # Skills
        if resume_data.get('skills'):
            doc.add_heading('SKILLS', level=2)
            skills_text = ' | '.join(resume_data['skills'])
            doc.add_paragraph(skills_text)
            doc.add_paragraph()
        
        # Experience
        if resume_data.get('experience'):
            doc.add_heading('EXPERIENCE', level=2)
            for job in resume_data['experience']:
                # Job title and company
                title_para = doc.add_paragraph()
                title_run = title_para.add_run(f"{job.get('title', '')} - {job.get('company', '')}")
                title_run.bold = True
                title_run.font.size = Pt(11)
                
                # Dates
                dates_para = doc.add_paragraph(f"{job.get('start_date', '')} to {job.get('end_date', '')}")
                dates_para.runs[0].font.size = Pt(10)
                dates_para.runs[0].font.italic = True
                
                # Description
                if job.get('description'):
                    desc_para = doc.add_paragraph(job['description'], style='List Bullet')
                
                doc.add_paragraph()
        
        # Education
        if resume_data.get('education'):
            doc.add_heading('EDUCATION', level=2)
            for edu in resume_data['education']:
                edu_para = doc.add_paragraph()
                edu_run = edu_para.add_run(f"{edu.get('degree', '')} in {edu.get('field', '')}")
                edu_run.bold = True
                
                school_para = doc.add_paragraph(edu.get('school', ''))
                school_para.runs[0].font.size = Pt(10)
                
                if edu.get('graduation_date'):
                    date_para = doc.add_paragraph(f"Graduated: {edu.get('graduation_date', '')}")
                    date_para.runs[0].font.size = Pt(10)
                
                doc.add_paragraph()
        
        # Certifications
        if resume_data.get('certifications'):
            doc.add_heading('CERTIFICATIONS', level=2)
            for cert in resume_data['certifications']:
                doc.add_paragraph(cert, style='List Bullet')
        
        # Determine output path
        if output_dir is None:
            output_dir = str(Path(__file__).parent.parent.parent)  # Project root
        
        os.makedirs(output_dir, exist_ok=True)
        
        # Generate filename with timestamp
        from datetime import datetime
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        name_slug = resume_data.get('name', 'resume').replace(' ', '_').lower()
        output_filename = f"tailored_resume_{name_slug}_{timestamp}.docx"
        output_path = os.path.join(output_dir, output_filename)
        
        # Save document
        doc.save(output_path)
        
        return {
            "status": "success",
            "message": f"Resume generated successfully",
            "file_path": output_path,
            "file_name": output_filename
        }
    
    except Exception as e:
        return {
            "status": "error",
            "message": f"Failed to generate resume: {str(e)}",
            "error": str(e)
        }
